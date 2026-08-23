from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Iterable

from docx import Document


PLACEHOLDER_PATTERN = re.compile(
    r"\{\{\s*([A-Za-z][A-Za-z0-9_]*)\s*\}\}"
)


@dataclass(frozen=True)
class TemplateValidationResult:
    placeholders: list[str]
    known_placeholders: list[str]
    unknown_placeholders: list[str]
    missing_required_fields: list[str]
    duplicate_placeholders: list[str]
    is_valid: bool


def extract_placeholders_from_text(text: str) -> list[str]:
    return [
        match.group(1).strip()
        for match in PLACEHOLDER_PATTERN.finditer(text or "")
    ]


def _iter_paragraph_texts(document: Document) -> Iterable[str]:
    for paragraph in document.paragraphs:
        yield paragraph.text

    # Tables are commonly used for certificate layouts.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text

    # Headers and footers often contain university names/logos and
    # occasionally dynamic placeholders such as {{issue_date}}.
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph.text

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph.text

        for paragraph in section.footer.paragraphs:
            yield paragraph.text

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph.text


def extract_docx_placeholders(file_bytes: bytes) -> list[str]:
    document = Document(io.BytesIO(file_bytes))

    placeholders: list[str] = []

    for text in _iter_paragraph_texts(document):
        placeholders.extend(
            extract_placeholders_from_text(text)
        )

    # Preserve first-seen order while removing duplicates.
    unique: list[str] = []
    seen: set[str] = set()

    for placeholder in placeholders:
        if placeholder not in seen:
            seen.add(placeholder)
            unique.append(placeholder)

    return unique


def validate_docx_template(
    *,
    file_bytes: bytes,
    configured_field_keys: Iterable[str],
    required_field_keys: Iterable[str],
    optional_system_keys: Iterable[str] | None = None,
) -> TemplateValidationResult:
    configured = {
        key.strip()
        for key in configured_field_keys
        if key and key.strip()
    }

    required = {
        key.strip()
        for key in required_field_keys
        if key and key.strip()
    }

    system_keys = {
        key.strip()
        for key in (
            optional_system_keys
            or {
                "issue_date",
                "certificate_number",
                "current_date",
            }
        )
        if key and key.strip()
    }

    placeholders = extract_docx_placeholders(
        file_bytes
    )

    placeholder_set = set(placeholders)
    known = placeholder_set & (configured | system_keys)
    unknown = placeholder_set - (configured | system_keys)

    missing = required - placeholder_set

    counts: dict[str, int] = {}

    for placeholder in placeholders:
        counts[placeholder] = counts.get(placeholder, 0) + 1

    duplicate_placeholders = sorted(
        placeholder
        for placeholder, count in counts.items()
        if count > 1
    )

    # Duplicate placeholders are allowed in a document. They are reported
    # for visibility, but they do not make the template invalid.
    is_valid = not unknown and not missing

    return TemplateValidationResult(
        placeholders=placeholders,
        known_placeholders=sorted(known),
        unknown_placeholders=sorted(unknown),
        missing_required_fields=sorted(missing),
        duplicate_placeholders=duplicate_placeholders,
        is_valid=is_valid,
    )